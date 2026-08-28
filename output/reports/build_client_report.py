from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = r"D:\Rapidsai_all_projects\Frontend_Crypto_project\output\reports\Client_Update_Report_August_24_2026.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)


def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(31, 78, 121)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Client Progress Update")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Daily Implementation Summary | August 24, 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(89, 89, 89)

    add_heading(doc, "Overview")
    add_body(
        doc,
        "Today focused on improving the scanner foundation and delivering the EMA screening experience requested for client review. "
        "The work completed today strengthens result reliability, improves alignment with TradingView checks, and adds a clearer user experience for EMA-based screening.",
    )

    add_heading(doc, "Completed Today")
    add_bullet(doc, "Improved scanner reliability by ensuring checks are based on completed candles.")
    add_bullet(doc, "Added EMA screening behavior that supports TradingView-style review for selected EMA levels.")
    add_bullet(doc, "Added support for the common EMA 20, 50, 100, and 200 setup used in TradingView.")
    add_bullet(doc, "Added client-facing EMA controls so users can select EMA periods, choose match behavior, and adjust the EMA condition range.")
    add_bullet(doc, "Added a quick preset for EMA 20/50/100/200 to make the TradingView comparison easier.")
    add_bullet(doc, "Added quality checks around the EMA behavior to confirm the expected passing cases.")

    add_heading(doc, "Testing Report Included")
    add_body(
        doc,
        "The supplied EMA testing report was reviewed and included in this update as client-facing evidence. "
        "It records passing EMA examples using TradingView as the reference view.",
    )

    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = "Table Grid"
    summary_rows = [
        ("Testing report", "EMA Indicator Testing, Passing Results Report"),
        ("Reference platform", "TradingView"),
        ("Test date", "24 August 2026"),
        ("Overall result", "PASS"),
        ("Evidence covered", "13 passing chart examples across 5m, 15m, 1h, and 1D timeframes"),
    ]
    for row_index, (label, value) in enumerate(summary_rows):
        cells = summary_table.rows[row_index].cells
        set_cell_text(cells[0], label, bold=True, color="1F4E79")
        set_cell_text(cells[1], value)
        set_cell_shading(cells[0], "F2F4F7")

    add_heading(doc, "Testing Evidence Summary")
    evidence = [
        ("AAMI", "1h", "EMA 50", "PASS"),
        ("AFYA", "1h", "EMA 20", "PASS"),
        ("HPP", "1h", "EMA 200", "PASS"),
        ("APAM", "1h", "EMA 20", "PASS"),
        ("ASIX", "1h", "EMA 20", "PASS"),
        ("ALG", "1D", "EMA 50", "PASS"),
        ("ADNT", "1D", "EMA 20", "PASS"),
        ("BOF", "1D", "EMA 100", "PASS"),
        ("BESS", "1D", "EMA 50", "PASS"),
        ("ACRV", "5m", "EMA 20", "PASS"),
        ("AMAT", "5m", "EMA 50 / 100", "PASS"),
        ("ASIX", "15m", "EMA 100", "PASS"),
        ("DVLT", "15m", "EMA 200", "PASS"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Symbol", "Timeframe", "Observed EMA", "Result"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, color="FFFFFF")
        set_cell_shading(cell, "1F4E79")
    for symbol, timeframe, observed, result in evidence:
        cells = table.add_row().cells
        set_cell_text(cells[0], symbol)
        set_cell_text(cells[1], timeframe)
        set_cell_text(cells[2], observed)
        set_cell_text(cells[3], result, bold=True, color="2E7D32")

    add_heading(doc, "Client Review Note")
    add_body(
        doc,
        "The EMA work completed today is ready for client review using the included passing examples. "
        "The next review can focus on confirming that the visible scanner results match the expected TradingView behavior for the selected EMA setup.",
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Client Progress Update | August 24, 2026")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_report()
